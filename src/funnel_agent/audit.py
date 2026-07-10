"""Full digital-marketing AUDIT for one prospect — a detailed, agency-grade report assembled from
the enrichment we already hold (DataForSEO SEO money-keywords, Google Ads Transparency creatives,
competitor gap analysis, backlinks, GEO/AEO technical checks) + the business firmographics, plus a
synthesised executive summary and a Traffic-Radius recommendation with a quantified opportunity.

`assemble_audit(pool, domain)` returns a rich, render-agnostic model. `render_audit_html(model)`
turns it into a premium, self-contained HTML report (used on-screen, as the PDF source, and for the
DOCX). No network calls here — it reads what enrichment already stored (audits are fetched elsewhere).
"""
from __future__ import annotations

import base64
import os
from datetime import date, datetime

_LOGO_CACHE: dict = {}


def _logo_datauri() -> str:
    """The real Traffic Radius logo as a data URI (embedded so the report is self-contained)."""
    if "u" not in _LOGO_CACHE:
        try:
            p = os.path.join(os.path.dirname(__file__), "dashboard", "static", "logo.png")
            b = open(p, "rb").read()
            _LOGO_CACHE["u"] = "data:image/png;base64," + base64.b64encode(b).decode()
        except Exception:
            _LOGO_CACHE["u"] = ""
    return _LOGO_CACHE["u"]


def embed_ad_images(model: dict, *, limit: int = 6, timeout: float = 15.0) -> dict:
    """Download each ad's Google-hosted preview image and inline it as a data URI so the actual ad
    creatives render in a self-contained report (Artifact/PDF/DOCX). Network side-effect; best-effort.
    Mutates + returns the model. Also stamps the embedded logo."""
    import httpx
    model.setdefault("_logo", _logo_datauri())
    creatives = ((model.get("ads") or {}).get("creatives")) or []
    todo = [c for c in creatives if c.get("preview_image")][:limit]
    if not todo:
        return model
    try:
        cli = httpx.Client(timeout=timeout, follow_redirects=True, headers={"User-Agent": "Mozilla/5.0"})
    except Exception:
        return model
    for c in todo:
        try:
            r = cli.get(c["preview_image"])
            ct = r.headers.get("content-type", "")
            if r.status_code == 200 and "image" in ct and len(r.content) < 900_000:
                c["image_data"] = f"data:{ct};base64," + base64.b64encode(r.content).decode()
        except Exception:
            continue
    cli.close()
    return model


# ---------- formatting ----------
def _n(v, d=0):
    try:
        return f"{float(v):,.{d}f}" if v is not None else "—"
    except (TypeError, ValueError):
        return "—"


def _money(v, d=0):
    try:
        return f"${float(v):,.{d}f}" if v is not None else "—"
    except (TypeError, ValueError):
        return "—"


def _pos(p):
    try:
        return f"#{int(p)}"
    except (TypeError, ValueError):
        return "—"


_ORGANIC_CTR = {1: .28, 2: .15, 3: .10, 4: .07, 5: .05, 6: .04, 7: .03, 8: .025, 9: .02, 10: .018}


def _ctr(pos) -> float:
    """Industry-average organic CTR by SERP position (mirrors dataforseo._ctr) — used to ESTIMATE
    traffic from search volume."""
    try:
        p = int(pos)
    except (TypeError, ValueError):
        return 0.0
    if p <= 10:
        return _ORGANIC_CTR.get(p, .018)
    if p <= 20:
        return .010
    if p <= 30:
        return .005
    return .002


def _get(d, *path, default=None):
    cur = d
    for k in path:
        if not isinstance(cur, dict):
            return default
        cur = cur.get(k)
    return cur if cur is not None else default


# ---------- assembler ----------
def assemble_audit(pool, domain: str, *, avg_ticket: float | None = None) -> dict | None:
    """Read every enrichment source for `domain` and assemble the audit model. `avg_ticket` (the
    prospect's average sale/project value) turns the SEO upside into a REVENUE opportunity. Returns
    None when the domain has no enrichment row at all."""
    from psycopg.rows import dict_row
    domain = (domain or "").strip().lower().replace("https://", "").replace("http://", "").strip("/")
    if not domain:
        return None
    with pool.connection() as conn, conn.cursor(row_factory=dict_row) as cur:
        cur.execute("SELECT domain, semrush, apollo, website, business_intel, whois, dataforseo, fetched_at "
                    "FROM enrichment WHERE domain=%s", (domain,))
        enr = cur.fetchone()
        cur.execute("SELECT company_name, industry, sub_industry, revenue_musd, employees, suburb, state "
                    "FROM companies WHERE domain=%s ORDER BY revenue_musd DESC NULLS LAST LIMIT 1", (domain,))
        comp = cur.fetchone()
    if not enr and not comp:
        return None
    enr = enr or {}
    df = enr.get("dataforseo") or {}
    web = enr.get("website") or {}
    ca = df.get("competitor_audit") or {}
    seo = df.get("audit") or {}
    ads = df.get("ads") or {}
    rank = df.get("rank") or {}
    org = rank.get("organic") or {}
    paid = rank.get("paid") or {}
    bi = enr.get("business_intel") or {}

    name = (comp or {}).get("company_name") or (bi.get("name") if isinstance(bi, dict) else None) or domain
    location = ", ".join(x for x in [(comp or {}).get("suburb"), (comp or {}).get("state")] if x) or None

    # brand tokens (used to strip the prospect's OWN name from every opportunity table/figure)
    try:
        from .enrichment.dataforseo import brand_tokens, is_branded as _is_branded
        _brands = brand_tokens(domain, name)
    except Exception:
        _brands = set()
        def _is_branded(kw, b):
            return False

    # Strip the prospect's OWN brand terms from the money-keyword & proof tables BEFORE any value is summed
    # — a business ranking for its own name isn't an "opportunity", and a high-CPC brand term (e.g. "clover"
    # at $19 CPC) otherwise dominates the table with an inflated, misleading $-value that kills credibility.
    def _debrand(rows):
        return [r for r in (rows or []) if not _is_branded(r.get("keyword"), _brands)]
    if isinstance(seo, dict):
        if seo.get("money_keywords"):
            seo["money_keywords"] = _debrand(seo["money_keywords"])
        if seo.get("proof_winning"):
            seo["proof_winning"] = _debrand(seo["proof_winning"])

    # ---- quantified opportunity ----
    seo_tot = seo.get("totals") or {}
    est_org_traffic = seo_tot.get("est_organic_traffic")
    est_org_value = seo_tot.get("est_traffic_value")
    quickwin_traffic = seo_tot.get("quickwin_upside_traffic")
    # $ value competitors capture on keywords the prospect doesn't rank for (the gap)
    gap_value = sum((k.get("money_value") or 0) for k in (ca.get("keyword_gap") or []))
    outranked_value = sum((k.get("money_value") or 0) for k in (ca.get("outranked") or []))
    # quick-win $ upside (reaching page 1 on the money keywords) — brand terms already removed
    quickwin_value = sum((k.get("upside_value") or 0) for k in (seo.get("money_keywords") or []))

    running_ads = bool(ads.get("running_ads")) or bool(df.get("running_google_ads"))
    geo = ca.get("geo_aeo") or {}

    # ---- keyword universe → clusters → funnel ----
    universe = build_keyword_universe(df.get("ranked_kw") or [], ca.get("keyword_gap") or [], brands=_brands)

    # ---- revenue opportunity model (avg-ticket driven) ----
    # Only COMMERCIAL-intent traffic converts to sales, so the addressable base is the quick-win
    # money keywords (page-2 upside) + the not-yet-ranked MOFU/BOFU gap — NOT informational (TOFU)
    # traffic. This keeps the revenue estimate defensible instead of inflating it with awareness clicks.
    _fn = universe.get("funnel") or {}
    commercial_gap = ((_fn.get("MOFU") or {}).get("traffic") or 0) + ((_fn.get("BOFU") or {}).get("traffic") or 0)
    addressable_traffic = int((quickwin_traffic or 0) + commercial_gap)
    VISITOR_TO_LEAD, LEAD_TO_SALE = 0.03, 0.25   # conservative service-business funnel assumptions
    rev_monthly = round(addressable_traffic * VISITOR_TO_LEAD * LEAD_TO_SALE * avg_ticket) if avg_ticket else None
    revenue = {
        "avg_ticket": avg_ticket,
        "addressable_traffic": addressable_traffic,
        "visitor_to_lead": VISITOR_TO_LEAD,
        "lead_to_sale": LEAD_TO_SALE,
        "leads_per_mo": round(addressable_traffic * VISITOR_TO_LEAD, 1),
        "sales_per_mo": round(addressable_traffic * VISITOR_TO_LEAD * LEAD_TO_SALE, 1),
        "monthly": rev_monthly,
        "annual": rev_monthly * 12 if rev_monthly else None,
    }

    ads_m = _ads_model(ads, paid)
    # ---- competitor scoreboard + share of voice ----
    comps = ca.get("competitors") or []
    comp_traffic = sum((c.get("est_traffic") or 0) for c in comps)
    our_traffic = est_org_traffic or org.get("etv") or 0
    sov = round(100 * our_traffic / max(1, our_traffic + comp_traffic), 1) if (our_traffic or comp_traffic) else None
    # ---- REALISTICALLY CAPTURABLE gap value (fix ~20x inflation: capture-CTR × CPC, not full volume) ----
    for k in (ca.get("keyword_gap") or []):
        k["cap_value"] = round((k.get("est_capture_traffic") or 0) * (k.get("cpc") or 0))
    for k in (ca.get("outranked") or []):
        cap_t = k.get("est_capture_traffic") or round((k.get("volume") or 0) * _ctr(5))
        k["cap_value"] = round(cap_t * (k.get("cpc") or 0))
    gap_capturable = sum(k.get("cap_value") or 0 for k in (ca.get("keyword_gap") or []))
    health = _health_scores(est_org_value, seo_tot, ads_m, sov, geo, len(comps))
    diagnosis = _diagnosis(name, running_ads, ads_m, quickwin_value, gap_capturable, sov, geo, comps)

    findings = _exec_findings(name, running_ads, ads, est_org_value, quickwin_value, gap_capturable,
                              ca, seo, org, geo)
    recommendation = _recommendation(name, running_ads, quickwin_value, gap_capturable, ca, seo, geo, web)

    return {
        "domain": domain,
        "name": name,
        "generated": (seo.get("generated_at") or ca.get("generated_at") or df.get("fetched_at") or ""),
        "business": {
            "industry": (comp or {}).get("industry"),
            "sub_industry": (comp or {}).get("sub_industry"),
            "revenue_musd": (comp or {}).get("revenue_musd"),
            "employees": (comp or {}).get("employees"),
            "location": location,
            "website": web.get("final_url") or f"https://{domain}",
        },
        "opportunity": {
            "est_org_traffic": est_org_traffic,
            "est_org_value": est_org_value,
            "quickwin_traffic": quickwin_traffic,
            "quickwin_value": quickwin_value,
            "gap_value": gap_value,
            "gap_capturable": gap_capturable,
            "outranked_value": outranked_value,
            "org_keywords": org.get("count"),
            "paid_keywords": paid.get("count"),
        },
        "health": health,
        "sov": sov,
        "our_traffic": our_traffic,
        "diagnosis": diagnosis,
        "ads": ads_m,
        "seo": {
            "totals": seo_tot,
            "money_keywords": seo.get("money_keywords") or [],
            "proof_winning": seo.get("proof_winning") or [],
            "growth": seo.get("growth") or [],
            "assumption": seo.get("assumption"),
            "keyword_count_total": seo.get("keyword_count_total") or org.get("count"),
        },
        "competitors": ca.get("competitors") or [],
        "keyword_gap": ca.get("keyword_gap") or [],
        "outranked": ca.get("outranked") or [],
        "content_gap": ca.get("content_gap") or [],
        "quick_wins": ca.get("quick_wins") or [],
        "growth_plan": ca.get("growth_plan") or [],
        "backlink_gap": ca.get("backlink_gap") or {},
        "geo_aeo": geo,
        "tech": {
            "trackers": web.get("trackers") or {},
            "socials": web.get("socials") or {},
            "https": (web.get("final_url") or "").startswith("https"),
            "final_url": web.get("final_url"),
        },
        "universe": universe,
        "revenue": revenue,
        "findings": findings,
        "recommendation": recommendation,
        "assumptions": ca.get("assumptions") or ([seo.get("assumption")] if seo.get("assumption") else []),
    }


def _ads_model(ads: dict, paid: dict) -> dict:
    """Rich Google-Ads model from the Transparency Center data (the authoritative source — DataForSEO
    Labs 'paid' keyword metrics under-report and are often 0 even for confirmed advertisers)."""
    items = ads.get("items") or []
    def _d(s):
        try:
            return datetime.fromisoformat(str(s).replace(" +00:00", "+00:00").replace(" ", "T", 1)[:19])
        except Exception:
            return None
    firsts = [d for d in (_d(c.get("first_shown")) for c in items) if d]
    lasts = [d for d in (_d(c.get("last_shown")) for c in items) if d]
    since = min(firsts).date().isoformat() if firsts else None
    latest = max(lasts).date().isoformat() if lasts else ads.get("last_shown")
    years = round((max(lasts) - min(firsts)).days / 365, 1) if (firsts and lasts) else None
    fmts = ads.get("formats") or {}
    # normalise format labels
    fmt_labels = {"text": "Search (text)", "image": "Display (image)", "video": "Video", "unknown": "Other"}
    formats = {fmt_labels.get(k, k.title()): v for k, v in fmts.items()}
    return {
        "running": bool(ads.get("running_ads")),
        "count": ads.get("count") or len(items),
        "verified": bool(ads.get("verified")),
        "last_shown": latest,
        "since": since,
        "years_active": years,
        "formats": formats,
        "advertisers": ads.get("advertisers") or [],
        "creatives": items,   # keep ALL; renderer picks the ones with images
        "paid_est_spend": (paid or {}).get("traffic_cost") or None,   # often 0/unreliable — shown only if >0
        "paid_keywords": (paid or {}).get("count") or None,
    }


def _exec_findings(name, running_ads, ads, est_org_value, quickwin_value, gap_value, ca, seo, org, geo):
    """3–5 punchy, quantified headline findings for the executive summary."""
    out = []
    if running_ads:
        out.append({"kind": "spend", "text": (
            f"{name} is actively running Google Ads — {ads.get('count') or 'multiple'} live creatives "
            f"confirmed in Google's Ads Transparency Center (last seen {ads.get('last_shown') or 'recently'}). "
            "They are already paying for traffic, which makes their SEO gaps expensive.")})
    if quickwin_value:
        out.append({"kind": "quickwin", "text": (
            f"There is roughly {_money(quickwin_value)}/mo of organic traffic value within reach from "
            "‘quick-win’ keywords already ranking on page 2 — a short push to page 1 captures it without ad spend.")})
    if gap_value:
        out.append({"kind": "gap", "text": (
            f"Competitors are capturing about {_money(gap_value)}/mo of search value on money keywords "
            f"{name} does not rank for at all — direct, quantified lost ground.")})
    comps = ca.get("competitors") or []
    if comps:
        top = comps[0]
        out.append({"kind": "competitor", "text": (
            f"Their strongest organic competitor, {top.get('domain')}, pulls an estimated "
            f"{_n(top.get('est_traffic'))} organic visits/mo — the benchmark to close.")})
    if geo and geo.get("score") is not None and geo.get("score") < 60:
        out.append({"kind": "tech", "text": (
            f"The site scores only {geo.get('score')}/100 on AI-search / answer-engine readiness — "
            "it risks being invisible as buyers shift to AI-assisted search.")})
    if not out and est_org_value:
        out.append({"kind": "seo", "text": (
            f"The site earns an estimated {_money(est_org_value)}/mo in organic traffic value today — "
            "a base we can materially grow.")})
    return out[:5]


def _recommendation(name, running_ads, quickwin_value, gap_value, ca, seo, geo, web):
    """The Traffic Radius recommendation: prioritised plan + quantified upside."""
    steps = []
    if seo.get("money_keywords"):
        steps.append(("Win the quick wins", (
            "Optimise the pages already ranking on page 2 for high-value, non-branded money keywords "
            "(on-page, internal linking, content depth) to move them onto page 1 — the fastest ROI.")))
    if ca.get("keyword_gap"):
        steps.append(("Close the competitor keyword gap", (
            "Build targeted content + landing pages for the money keywords competitors rank for and "
            f"{name} doesn’t, capturing demand that currently flows to rivals.")))
    if running_ads:
        steps.append(("Tighten paid + organic together", (
            "They already spend on Google Ads — align the campaigns with the SEO plan so paid covers "
            "the gaps while organic compounds, cutting cost-per-lead over time.")))
    if geo and geo.get("score") is not None and geo.get("score") < 70:
        steps.append(("Fix AI-search / answer-engine readiness", (
            "Add the structured data, clear entity signals and answer-formatted content that make the "
            "site quotable by Google AI Overviews and assistants.")))
    if ca.get("backlink_gap"):
        steps.append(("Build authority", (
            "A focused digital-PR + link-building push to close the referring-domain gap behind competitors.")))
    total_upside = (quickwin_value or 0) + (gap_value or 0)
    summary = (
        f"{name} is already investing in demand (live Google Ads) but leaving material organic value on "
        f"the table. Between quick-win rankings and the competitor keyword gap, roughly {_money(total_upside)}/mo "
        "of search value is addressable. Traffic Radius would combine SEO + paid management to capture it.")
    return {"summary": summary, "steps": steps[:5], "total_upside": total_upside}


# ---------- premium HTML renderer (on-screen + PDF source) ----------
def _esc(s):
    s = "" if s is None else str(s)
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;"))


def _kw_rows(rows, cols):
    """cols = list of (header, key, fmt) — fmt in {kw,pos,int,money,cpc,pct}."""
    def cell(r, key, fmt):
        v = r.get(key)
        if fmt == "kw":
            u = r.get("url")
            return f'<td class="kw">{_esc(v)}</td>'
        if fmt == "pos":
            return f'<td class="num">{_pos(v)}</td>'
        if fmt == "int":
            return f'<td class="num">{_n(v)}</td>'
        if fmt == "money":
            return f'<td class="num">{_money(v)}</td>'
        if fmt == "cpc":
            return f'<td class="num">{_money(v,2)}</td>'
        return f'<td>{_esc(v)}</td>'
    head = "".join(f'<th class="{ "kw" if f=="kw" else "num"}">{_esc(h)}</th>' for h, _, f in cols)
    body = "".join("<tr>" + "".join(cell(r, k, f) for _, k, f in cols) + "</tr>" for r in rows)
    return f'<table class="kw-table"><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>'


_AUDIT_CSS = """
*{box-sizing:border-box}
:root{--ink:#0f172a;--ink2:#334155;--muted:#64748b;--faint:#94a3b8;--line:#e2e8f0;--line2:#cbd5e1;
  --brand:#2563eb;--brand2:#1d4ed8;--brand-soft:#eff6ff;--good:#16a34a;--good-soft:#f0fdf4;
  --warn:#d97706;--warn-soft:#fffbeb;--bad:#dc2626;--paper:#ffffff;--panel:#f8fafc}
html,body{margin:0;padding:0;background:#eef2f7;color:var(--ink);
  font:14px/1.6 "Segoe UI",Roboto,Helvetica,Arial,ui-sans-serif,system-ui;-webkit-font-smoothing:antialiased;font-variant-numeric:tabular-nums}
.page{max-width:900px;margin:0 auto;background:var(--paper)}
.sec{padding:30px 46px;border-top:1px solid var(--line)}
.sec:first-of-type{border-top:0}
h1,h2,h3{margin:0;letter-spacing:-.02em;color:var(--ink)}
.eyebrow{font-size:11px;font-weight:800;letter-spacing:.14em;text-transform:uppercase;color:var(--brand)}
.secnum{display:inline-flex;align-items:center;justify-content:center;width:26px;height:26px;border-radius:7px;
  background:var(--brand-soft);color:var(--brand);font-weight:800;font-size:13px;margin-right:10px}
.sec h2{font-size:21px;font-weight:800;display:flex;align-items:center;margin-bottom:4px}
.sec .lead{color:var(--muted);font-size:13.5px;margin:6px 0 18px}
/* cover */
.cover{padding:52px 46px 34px;background:linear-gradient(160deg,#0b1e46 0%,#12306e 55%,#1d4ed8 100%);color:#fff;position:relative;overflow:hidden}
.cover::after{content:"";position:absolute;right:-80px;top:-80px;width:320px;height:320px;border-radius:50%;
  background:radial-gradient(circle,rgba(255,255,255,.12),transparent 70%)}
.cover .brand{display:flex;align-items:center;gap:11px;font-weight:800;letter-spacing:.02em;font-size:15px;opacity:.98}
.cover .brand .logo{background:#fff;border-radius:8px;padding:6px 9px;display:inline-flex;align-items:center;justify-content:center;box-shadow:0 2px 8px rgba(0,0,0,.25)}
.cover .brand .logo img{height:22px;display:block}
.cover .kick{margin-top:34px;font-size:12px;font-weight:800;letter-spacing:.16em;text-transform:uppercase;color:#9dc0ff}
.cover h1{color:#fff;font-size:34px;font-weight:800;line-height:1.1;margin:8px 0 6px}
.cover .dom{color:#bcd3ff;font-size:15px;font-weight:600}
.cover .meta{margin-top:22px;display:flex;gap:26px;flex-wrap:wrap;color:#cfe0ff;font-size:12.5px}
.cover .meta b{color:#fff;display:block;font-size:15px;font-weight:800;margin-top:2px}
.headline{margin-top:28px;background:rgba(255,255,255,.10);border:1px solid rgba(255,255,255,.20);border-radius:14px;padding:20px 22px}
.headline .big{font-size:44px;font-weight:800;color:#fff;letter-spacing:-.03em;line-height:1}
.headline .big .per{font-size:19px;font-weight:600;color:#9dc0ff;letter-spacing:0;margin-left:4px}
.headline .cap{color:#cfe0ff;font-size:13px;margin-top:7px;line-height:1.45;max-width:520px}
/* scorecards */
.cards{display:flex;gap:14px;flex-wrap:wrap}
.card{flex:1 1 150px;border:1px solid var(--line);border-radius:12px;padding:15px 16px;background:var(--panel)}
.card .v{font-size:24px;font-weight:800;letter-spacing:-.02em;line-height:1.1}
.card .l{font-size:11px;font-weight:700;letter-spacing:.05em;text-transform:uppercase;color:var(--muted);margin-top:6px}
.card .s{font-size:11.5px;color:var(--faint);margin-top:3px}
.card.good{background:var(--good-soft);border-color:#bbf7d0}.card.good .v{color:var(--good)}
.card.brand{background:var(--brand-soft);border-color:#bfdbfe}.card.brand .v{color:var(--brand)}
/* findings */
.finding{display:flex;gap:12px;padding:12px 0;border-bottom:1px solid var(--line)}
.finding:last-child{border-bottom:0}
.finding .dot{flex:none;width:9px;height:9px;border-radius:50%;background:var(--brand);margin-top:7px}
.finding.spend .dot{background:var(--warn)}.finding.gap .dot{background:var(--bad)}.finding.quickwin .dot{background:var(--good)}
.finding .t{font-size:13.5px;color:var(--ink2)}
/* tables */
.kw-table{width:100%;border-collapse:collapse;font-size:12.5px;margin-top:8px}
.kw-table th{text-align:right;padding:8px 10px;border-bottom:2px solid var(--line2);color:var(--muted);
  font-size:10.5px;font-weight:800;letter-spacing:.05em;text-transform:uppercase}
.kw-table th.kw,.kw-table td.kw{text-align:left}
.kw-table td{padding:8px 10px;border-bottom:1px solid var(--line);vertical-align:top}
.kw-table td.kw{font-weight:600;color:var(--ink)}
.kw-table td.num{text-align:right;font-variant-numeric:tabular-nums;color:var(--ink2)}
.kw-table tbody tr:nth-child(even){background:var(--panel)}
/* competitor rows */
.comp{display:flex;align-items:center;gap:14px;padding:11px 0;border-bottom:1px solid var(--line)}
.comp:last-child{border-bottom:0}
.comp .rk{flex:none;width:24px;height:24px;border-radius:6px;background:var(--ink);color:#fff;font-weight:800;font-size:12px;display:flex;align-items:center;justify-content:center}
.comp .d{flex:1;font-weight:700}.comp .d small{display:block;color:var(--faint);font-weight:500;font-size:11.5px}
.comp .m{text-align:right;font-size:12.5px;color:var(--muted)}.comp .m b{display:block;color:var(--ink);font-size:15px}
/* creative gallery (real Google Ads Transparency images) */
.adgrid{display:grid;grid-template-columns:repeat(3,1fr);gap:12px;margin-top:12px}
.adcard{border:1px solid var(--line);border-radius:10px;overflow:hidden;background:#fff;display:flex;flex-direction:column;break-inside:avoid}
.adcard .imgwrap{background:#f1f5f9;display:flex;align-items:center;justify-content:center;min-height:130px;border-bottom:1px solid var(--line);padding:8px}
.adcard img{max-width:100%;max-height:210px;display:block;border-radius:4px}
.adcard .noimg{color:var(--faint);font-size:11px;text-align:center;padding:22px 8px}
.adcard .meta{padding:9px 11px}
.adcard .fmt{font-weight:800;color:var(--brand);font-size:10px;letter-spacing:.06em;text-transform:uppercase}
.adcard .dt{color:var(--faint);font-size:11px;margin-top:3px}
@media print{.adgrid{grid-template-columns:repeat(2,1fr)}}
/* recommendation */
.rec{background:var(--brand-soft);border:1px solid #bfdbfe;border-radius:12px;padding:18px 20px}
.rec .sum{font-size:14.5px;color:var(--ink);font-weight:600}
.step{display:flex;gap:12px;padding:11px 0;border-top:1px solid #dbeafe}
.step .n{flex:none;width:22px;height:22px;border-radius:50%;background:var(--brand);color:#fff;font-weight:800;font-size:12px;display:flex;align-items:center;justify-content:center;margin-top:1px}
.step .h{font-weight:800;font-size:13.5px}.step .d{color:var(--ink2);font-size:12.5px;margin-top:1px}
.geo{display:flex;align-items:center;gap:16px}
.geo .score{flex:none;width:74px;height:74px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-weight:800;font-size:22px;border:5px solid var(--good);color:var(--good)}
.geo .score.mid{border-color:var(--warn);color:var(--warn)}.geo .score.low{border-color:var(--bad);color:var(--bad)}
.checks{list-style:none;margin:8px 0 0;padding:0;columns:2;column-gap:26px;font-size:12.5px}
.checks li{padding:3px 0;break-inside:avoid}
.checks .ok{color:var(--good)}.checks .no{color:var(--bad)}
/* keyword funnel */
.kfunnel{margin-top:10px}
.kf-tier{margin:0 auto 8px;color:#fff;border-radius:8px;padding:12px 18px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:6px}
.kf-tier .kf-lab{font-weight:800;font-size:13.5px}
.kf-tier .kf-m{font-size:12px;opacity:.96;font-variant-numeric:tabular-nums;text-align:right}
.kf-tier.tofu{background:linear-gradient(90deg,#60a5fa,#3b82f6)}
.kf-tier.mofu{background:linear-gradient(90deg,#3b82f6,#2563eb)}
.kf-tier.bofu{background:linear-gradient(90deg,#1d4ed8,#1e3a8a)}
/* clusters */
.cltable{width:100%;border-collapse:collapse;font-size:12.5px;margin-top:10px}
.cltable th{text-align:right;padding:8px 10px;border-bottom:2px solid var(--line2);color:var(--muted);font-size:10.5px;font-weight:800;text-transform:uppercase;letter-spacing:.05em}
.cltable th.l,.cltable td.l{text-align:left}
.cltable td{padding:9px 10px;border-bottom:1px solid var(--line);vertical-align:middle}
.cltable td.num{text-align:right;font-variant-numeric:tabular-nums}
.cltable td.l{font-weight:700;color:var(--ink)}
.cltable tbody tr:nth-child(even){background:var(--panel)}
.covbar{display:inline-block;width:74px;height:7px;border-radius:4px;background:#e2e8f0;overflow:hidden;vertical-align:middle;margin-right:7px}
.covbar span{display:block;height:100%;background:var(--good)}
.fmix{display:inline-flex;gap:3px;align-items:flex-end;height:16px}
.fmix i{display:inline-block;width:14px;border-radius:2px}
.fmix .t{background:#93c5fd}.fmix .m{background:#3b82f6}.fmix .b{background:#1e3a8a}
/* revenue calculator */
.rev{background:linear-gradient(160deg,#0b1e46,#12306e);color:#fff;border-radius:16px;padding:24px 26px;margin-top:4px}
.rev .tkt{display:flex;align-items:center;gap:12px;flex-wrap:wrap;margin-bottom:20px;padding-bottom:18px;border-bottom:1px solid rgba(255,255,255,.12)}
.rev .tkt label{font-size:13px;color:#cfe0ff;font-weight:700}
.rev .tkt .box{background:rgba(255,255,255,.13);border:1px solid rgba(255,255,255,.32);border-radius:10px;padding:7px 13px;display:flex;align-items:center;gap:2px;font-size:19px;font-weight:800;transition:border-color .15s}
.rev .tkt .box:focus-within{border-color:#7ea6ff}
.rev .tkt .tkt-input{background:transparent;border:0;color:#fff;font:inherit;font-size:19px;font-weight:800;width:92px;outline:none}
.rev .tkt .tkthint{font-size:12px;color:#9db8ef;flex:1 1 200px;line-height:1.4}
.rev .chain{display:flex;gap:8px;flex-wrap:wrap;align-items:stretch}
.rev .arw{display:flex;align-items:center;color:#5f82c8;font-size:18px;font-weight:700}
.rev .s2{flex:1 1 108px;background:rgba(255,255,255,.07);border-radius:11px;padding:13px 15px}
.rev .s2 .v{font-size:21px;font-weight:800;letter-spacing:-.01em}.rev .s2 .l{font-size:11px;color:#bcd3ff;margin-top:4px;line-height:1.35}.rev .s2 .l span{color:#8fb0ff}
.rev .final{flex:1.5 1 188px;background:rgba(74,222,128,.15);border:1px solid rgba(74,222,128,.42);border-radius:11px;padding:13px 17px;display:flex;flex-direction:column;justify-content:center}
.rev .final .v{font-size:30px;font-weight:800;color:#86efac;letter-spacing:-.02em}.rev .final .l{font-size:12px;color:#bbf7d0;margin-top:2px}
.rev .ann{margin-top:15px;font-size:13.5px;color:#cfe0ff}.rev .ann b{color:#fff;font-size:18px}
.rev .note{margin-top:11px;font-size:11px;color:#7ea6ff;line-height:1.5}
/* CTA button */
.ctabtn{display:inline-block;margin-top:16px;background:#fff;color:#0b1e46;font-weight:800;font-size:14px;padding:12px 22px;border-radius:10px;letter-spacing:-.01em;text-decoration:none;cursor:pointer}
/* digital health scorecard — premium report card (grade + per-dimension score bars) */
.scorewrap{display:grid;grid-template-columns:212px 1fr;gap:26px;align-items:center;margin-top:4px}
@media(max-width:640px){.scorewrap{grid-template-columns:1fr;gap:18px}}
.grade{background:linear-gradient(160deg,#0b1e46,#12306e);border-radius:16px;padding:22px 18px;text-align:center;color:#fff}
.grade .gletter{font-family:Georgia,'Times New Roman',serif;font-size:92px;font-weight:700;line-height:.82;letter-spacing:-.02em}
.grade .gmeta{margin-top:12px}
.grade .gnum{font-size:26px;font-weight:800;color:#fff}.grade .gnum span{font-size:14px;font-weight:600;color:#8fb0ff}
.grade .gverdict{font-size:12px;font-weight:700;line-height:1.4;margin-top:6px}
.hbars{display:flex;flex-direction:column;gap:16px}
.hbar{display:grid;grid-template-columns:150px 1fr 34px;grid-template-rows:auto auto;gap:2px 12px;align-items:center}
.hbar .hbl{grid-row:1;grid-column:1;font-size:12.5px;font-weight:700;color:var(--ink)}
.hbar .htrack{grid-row:1;grid-column:2;position:relative;height:11px;background:var(--line2);border-radius:6px;overflow:visible}
.hbar .hfill{height:100%;border-radius:6px;transition:width .5s ease}
.hbar .htgt{position:absolute;top:-3px;height:17px;width:2px;background:var(--muted);opacity:.55;border-radius:2px}
.hbar .hsc{grid-row:1;grid-column:3;text-align:right;font-size:15px;font-weight:800;font-variant-numeric:tabular-nums}
.hbar .hnote{grid-row:2;grid-column:1 / span 3;font-size:11px;color:var(--muted);line-height:1.35}
/* diagnosis banner */
.diag{background:var(--brand-soft);border-left:4px solid var(--brand);border-radius:0 10px 10px 0;padding:14px 18px;font-size:13.5px;color:var(--ink);margin-top:16px;line-height:1.6}
/* cost of inaction */
.coi{display:flex;gap:12px;flex-wrap:wrap;margin-top:12px}
.coi .t{flex:1 1 130px;border:1px solid #fecaca;background:#fef2f2;border-radius:12px;padding:14px 16px}
.coi .t .v{font-size:22px;font-weight:800;color:var(--bad)}.coi .t .l{font-size:11px;font-weight:700;color:#b91c1c;text-transform:uppercase;letter-spacing:.04em;margin-top:4px}
/* competitor scoreboard + SoV */
.sovbar{display:flex;height:34px;border-radius:8px;overflow:hidden;border:1px solid var(--line);margin-top:10px}
.sovbar .seg{display:flex;align-items:center;justify-content:center;color:#fff;font-size:11px;font-weight:700;min-width:0;overflow:hidden;white-space:nowrap}
.scoreboard{width:100%;border-collapse:collapse;font-size:12.5px;margin-top:14px}
.scoreboard th{text-align:right;padding:8px 10px;border-bottom:2px solid var(--line2);color:var(--muted);font-size:10.5px;font-weight:800;text-transform:uppercase}
.scoreboard th.l,.scoreboard td.l{text-align:left}
.scoreboard td{padding:9px 10px;border-bottom:1px solid var(--line);text-align:right;font-variant-numeric:tabular-nums}
.scoreboard tr.us{background:var(--brand-soft)}
.scoreboard tr.us td{font-weight:800}.scoreboard tr.us td.l{color:var(--brand)}
/* battle cards */
.battle{display:grid;grid-template-columns:repeat(2,1fr);gap:12px;margin-top:12px}
.bcard{border:1px solid var(--line);border-radius:10px;padding:13px 15px;background:var(--panel)}
.bcard .h{font-weight:800;font-size:13.5px}.bcard .m{font-size:12px;color:var(--muted);margin-top:4px;line-height:1.5}.bcard .m b{color:var(--bad)}
@media print{.battle{grid-template-columns:repeat(2,1fr)}}
/* content gap — demand-ranked topic cards */
.cgstat{display:flex;gap:12px;flex-wrap:wrap;margin:0 0 16px}
.cgstat>div{flex:1 1 150px;background:var(--panel);border:1px solid var(--line);border-radius:12px;padding:14px 16px}
.cgstat .v{font-size:24px;font-weight:800;color:var(--ink);letter-spacing:-.01em}
.cgstat .l{font-size:11px;color:var(--muted);margin-top:3px;line-height:1.35}
.cggrid{display:grid;grid-template-columns:1fr 1fr;gap:12px}
@media(max-width:640px){.cggrid{grid-template-columns:1fr}}
.cgcard{border:1px solid var(--line);border-radius:12px;padding:14px 16px;background:var(--paper)}
.cgtop{display:flex;justify-content:space-between;align-items:baseline;gap:10px}
.cgt{font-weight:800;font-size:14.5px;letter-spacing:-.01em}
.cgv{font-size:16px;font-weight:800;color:var(--brand);white-space:nowrap}.cgv span{font-size:11px;font-weight:600;color:var(--muted)}
.cgbar{height:6px;background:var(--line2);border-radius:4px;overflow:hidden;margin:9px 0}
.cgbar span{display:block;height:100%;background:linear-gradient(90deg,var(--brand),#60a5fa);border-radius:4px}
.cgmeta{font-size:11.5px;color:var(--muted)}
.cgchips{display:flex;gap:6px;flex-wrap:wrap;margin-top:9px}
.cgchip{font-size:10.5px;background:var(--brand-soft);color:var(--brand2);border-radius:5px;padding:3px 7px;font-weight:600}
/* 90-day plan */
.phase{border-left:3px solid var(--brand);padding:2px 0 2px 16px;margin-bottom:16px}
.phase .ph{font-size:11.5px;font-weight:800;color:var(--brand);text-transform:uppercase;letter-spacing:.05em}
.pitem{padding:8px 0;border-bottom:1px solid var(--line)}
.pitem:last-child{border-bottom:0}
.pitem .t{font-weight:700;font-size:13px}.pitem .d{font-size:12px;color:var(--muted);margin-top:2px}
.chip{display:inline-block;font-size:9.5px;font-weight:800;text-transform:uppercase;letter-spacing:.04em;padding:2px 7px;border-radius:20px;margin-left:6px;vertical-align:middle}
.chip.hi{background:#dcfce7;color:#15803d}.chip.md{background:#fef3c7;color:#b45309}.chip.lo{background:#e0e7ff;color:#4338ca}
/* CTA */
.cta{background:linear-gradient(135deg,#1d4ed8,#12306e);color:#fff;border-radius:14px;padding:26px 30px;text-align:center;margin-top:6px}
.cta h3{font-size:22px;font-weight:800;color:#fff;margin:0 0 8px}
.cta p{color:#cfe0ff;font-size:14px;max-width:580px;margin:0 auto 6px;line-height:1.55}
.cta .rr{margin-top:14px;font-size:12.5px;color:#9dc0ff}
.assump{font-size:11px;color:var(--faint);line-height:1.5}
.footer{padding:22px 46px;background:#0b1e46;color:#9dc0ff;font-size:11.5px}
.footer b{color:#fff}
@media print{body{background:#fff}.page{max-width:none;margin:0}.sec{page-break-inside:avoid}}
"""


def render_audit_html(m: dict, *, standalone: bool = True) -> str:
    """Premium, self-contained HTML report for the audit model."""
    b = m["business"]; op = m["opportunity"]; ads = m["ads"]; seo = m["seo"]
    upside = m["recommendation"].get("total_upside") or 0
    gen = (m.get("generated") or "")[:10] or date.today().isoformat()

    # Revenue-opportunity base — every figure derived from the average ticket is tagged data-rev so a
    # single ticket input recomputes ALL of them live, everywhere in the report.
    rev = m.get("revenue") or {}
    r_traf = rev.get("addressable_traffic") or 0
    r_v2l = rev.get("visitor_to_lead") or 0.03
    r_l2s = rev.get("lead_to_sale") or 0.25
    r_ticket = int(rev.get("avg_ticket") or 2000)
    r_leads = r_traf * r_v2l
    r_sales = r_leads * r_l2s
    r_mo = round(r_sales * r_ticket)

    def rev_num(kind, val, money=True, dec=0):
        return f'<span data-rev="{kind}">{("$" if money else "")}{_n(val, dec)}</span>'

    # cover meta
    meta = []
    if b.get("industry"):
        meta.append(("Industry", b["industry"]))
    if b.get("location"):
        meta.append(("Location", b["location"]))
    if b.get("revenue_musd"):
        meta.append(("Revenue", f"${_n(b['revenue_musd'])}M"))
    if op.get("org_keywords"):
        meta.append(("Organic keywords", _n(op["org_keywords"])))
    meta_html = "".join(f'<div>{_esc(l)}<b>{_esc(v)}</b></div>' for l, v in meta)

    findings = "".join(
        f'<div class="finding {f.get("kind","")}"><div class="dot"></div><div class="t">{_esc(f["text"])}</div></div>'
        for f in m["findings"])

    # scorecards
    cards = []
    cards.append(("good", rev_num("mo", r_mo), "New revenue / mo", "at their average ticket — editable"))
    if op.get("est_org_value") is not None:
        cards.append(("brand", _money(op["est_org_value"]), "Organic value / mo", "estimated traffic value today"))
    if ads.get("running"):
        cards.append(("", (str(ads.get("count") or "Yes")), "Live Google Ads", f"confirmed · last seen {ads.get('last_shown') or '—'}"))
    if op.get("org_keywords"):
        cards.append(("", _n(op["org_keywords"]), "Ranked keywords", f"~{_n(op.get('est_org_traffic'))} visits/mo est."))
    cards_html = "".join(f'<div class="card {c}"><div class="v">{v}</div><div class="l">{l}</div><div class="s">{s}</div></div>' for c, v, l, s in cards)

    # sections are stored by key, then emitted in an explicit LOGICAL order + numbered at the end.
    S = {}
    def sec(key, title, lead, body, num=True):
        S[key] = (title, lead, body, num)

    # 0 Digital health scorecard — a premium "report card": overall grade + per-dimension score bars
    # with the GAP (unfilled portion) called out, so the prospect sees exactly what they're missing.
    h = m.get("health") or {}
    if h:
        ov = h.get("overall") or 0
        grade = "A" if ov >= 80 else "B" if ov >= 65 else "C" if ov >= 50 else "D" if ov >= 35 else "F"
        verdict = ("Strong foundations" if ov >= 65 else
                   "Underperforming — real gaps to close" if ov >= 45 else
                   "Significant gaps are holding this business back online")
        gcol = "#16a34a" if ov >= 65 else "#d97706" if ov >= 45 else "#dc2626"
        DIMS = [("seo", "Organic visibility", "how much of your market's search you actually capture"),
                ("ads", "Paid presence", "whether you're competing in the paid results"),
                ("competitive", "Competitive position", "your share of search vs your rivals"),
                ("technical", "AI-search readiness", "whether AI & answer engines can surface you")]
        bars = ""
        for k2, lab, note in DIMS:
            sc = int(h.get(k2) or 0)
            bcol = "#16a34a" if sc >= 65 else "#d97706" if sc >= 45 else "#dc2626"
            bars += (f'<div class="hbar"><div class="hbl">{_esc(lab)}</div>'
                     f'<div class="htrack"><div class="hfill" style="width:{sc}%;background:{bcol}"></div>'
                     f'<span class="htgt" style="left:80%"></span></div>'
                     f'<div class="hsc" style="color:{bcol}">{sc}</div>'
                     f'<div class="hnote">{_esc(note)}</div></div>')
        grade_card = (f'<div class="grade"><div class="gletter" style="color:{gcol}">{grade}</div>'
                      f'<div class="gmeta"><div class="gnum">{ov}<span>/100</span></div>'
                      f'<div class="gverdict" style="color:{gcol}">{_esc(verdict)}</div></div></div>')
        diag_html = f'<div class="diag">{_esc(m.get("diagnosis"))}</div>' if m.get("diagnosis") else ""
        sec("health", "Digital health scorecard",
            "An honest, at-a-glance grade of the business online today — the red gaps are the ground your competitors are taking.",
            f'<div class="scorewrap">{grade_card}<div class="hbars">{bars}</div></div>{diag_html}')

    # 1 Executive summary
    sec("exec", "Executive summary", "The headline findings and the size of the opportunity.",
        f'<div class="cards" style="margin-bottom:18px">{cards_html}</div>{findings}')

    # 2 Revenue opportunity (interactive — driven by the prospect's average ticket).
    # The ticket input is the ONE control that recomputes every data-rev figure across the whole report.
    rev_body = (
        f'<div class="rev">'
        f'<div class="tkt"><label>Their average sale / project value</label>'
        f'<span class="box">$<input class="tkt-input" type="number" min="0" step="100" value="{r_ticket}" oninput="recalcRev(this)"></span>'
        f'<span class="tkthint">set their real number — every revenue figure in this report updates instantly</span></div>'
        f'<div class="chain">'
        f'<div class="s2"><div class="v">{_n(r_traf)}</div><div class="l">extra visits / mo<br><span>addressable search</span></div></div>'
        f'<div class="arw">→</div>'
        f'<div class="s2"><div class="v">{rev_num("leads", r_leads, money=False, dec=1)}</div><div class="l">enquiries / mo<br><span>× {int(r_v2l*100)}%</span></div></div>'
        f'<div class="arw">→</div>'
        f'<div class="s2"><div class="v">{rev_num("sales", r_sales, money=False, dec=1)}</div><div class="l">new sales / mo<br><span>× {int(r_l2s*100)}%</span></div></div>'
        f'<div class="arw">→</div>'
        f'<div class="final"><div class="v">{rev_num("mo", r_mo)}</div><div class="l">new revenue / month</div></div>'
        f'</div>'
        f'<div class="ann">≈ <b>{rev_num("yr", r_mo*12)}</b> / year in revenue they are currently missing — from the organic search gap alone.</div>'
        f'<div class="note">Assumes {int(r_v2l*100)}% visitor→enquiry and {int(r_l2s*100)}% enquiry→sale — deliberately conservative service-business rates. Adjust the ticket above to their real average sale value.</div>'
        f'</div>')
    # cost of inaction — ticket-independent capturable search value slipping away over time
    coi_mo = (op.get("quickwin_value") or 0) + (op.get("gap_capturable") or 0)
    coi = ""
    if coi_mo:
        coi = ('<div class="lead" style="margin-top:20px">The cost of waiting — realistically capturable search value going uncaptured every…</div>'
               f'<div class="coi"><div class="t"><div class="v">{_money(coi_mo)}</div><div class="l">a month</div></div>'
               f'<div class="t"><div class="v">{_money(coi_mo*3)}</div><div class="l">a quarter</div></div>'
               f'<div class="t"><div class="v">{_money(coi_mo*12)}</div><div class="l">a year</div></div></div>')
    sec("revenue", "The revenue opportunity",
        "What the missed search traffic is worth in the prospect's own revenue terms — not just SEO metrics.", rev_body + coi)

    # 2 Google Ads audit
    if ads.get("running") or ads.get("count"):
        fmts = " · ".join(f"{k}: {v}" for k, v in (ads.get("formats") or {}).items())
        yrs = ads.get("years_active")
        since = ads.get("since")
        cards2 = ['<div class="card good"><div class="v">✓ Confirmed</div><div class="l">Running Google Ads</div>'
                  '<div class="s">via Google Ads Transparency Center</div></div>',
                  f'<div class="card"><div class="v">{_n(ads.get("count"))}</div><div class="l">Live creatives</div>'
                  f'<div class="s">{_esc(fmts) or "—"}</div></div>']
        if since:
            cards2.append(f'<div class="card brand"><div class="v">{_esc(since[:4])}</div><div class="l">Advertising since</div>'
                          f'<div class="s">{_n(yrs)+" yrs of sustained spend" if yrs else "ongoing"}</div></div>')
        if ads.get("verified"):
            cards2.append('<div class="card"><div class="v">Verified</div><div class="l">Advertiser identity</div>'
                          '<div class="s">Google-verified account</div></div>')
        # ad gallery — real Transparency Center creatives (embedded image_data if present)
        gallery = []
        for c in (ads.get("creatives") or []):
            img = c.get("image_data")
            if not img and not c.get("preview_image"):
                continue
            src = img or c.get("preview_image")
            fmt = {"text": "Search (text) ad", "image": "Display ad", "video": "Video ad"}.get(c.get("format"), (c.get("format") or "ad").title())
            inner = (f'<div class="imgwrap"><img src="{_esc(src)}" alt="Google Ad creative"/></div>' if src
                     else '<div class="imgwrap"><div class="noimg">(no preview available)</div></div>')
            gallery.append(f'<div class="adcard">{inner}<div class="meta"><div class="fmt">{_esc(fmt)}</div>'
                           f'<div class="dt">first shown {_esc((c.get("first_shown") or "?")[:10])} · last {_esc((c.get("last_shown") or "?")[:10])}</div></div></div>')
            if len(gallery) >= 6:
                break
        gal = f'<div class="adgrid">{"".join(gallery)}</div>' if gallery else ""
        gal_cap = ('<div class="lead" style="margin-top:16px">The actual ads they are running right now (Google Ads Transparency Center):</div>' if gallery else "")
        # insight line
        insight = ""
        if yrs and yrs >= 1:
            insight = (f'<div class="finding spend" style="border:0;padding-top:14px"><div class="dot"></div><div class="t">'
                       f'Google’s Transparency Center shows ads for this advertiser spanning <b>{_esc(since or "")} → today</b> '
                       f'(~{_n(yrs)} years, {_n(ads.get("count"))} creatives; {_esc(fmts)}) — a clear, ongoing paid-search commitment. '
                       f'That budget works far harder once the organic and landing-page foundation below is fixed.</div></div>')
        body = f'<div class="cards" style="margin-bottom:6px">{"".join(cards2)}</div>{insight}{gal_cap}{gal}'
        sec("gads", "Google Ads audit",
            "Independently verified from Google's public Ads Transparency Center — the authoritative record of the ads they actually run.", body)

    # 3 Organic SEO health
    if seo.get("money_keywords") or seo.get("totals"):
        t = seo.get("totals") or {}
        cards2 = (f'<div class="cards" style="margin-bottom:14px">'
                  f'<div class="card"><div class="v">{_n(t.get("keywords") or seo.get("keyword_count_total"))}</div><div class="l">Ranked keywords</div></div>'
                  f'<div class="card"><div class="v">{_n(t.get("est_organic_traffic"))}</div><div class="l">Est. visits / mo</div></div>'
                  f'<div class="card brand"><div class="v">{_money(t.get("est_traffic_value"))}</div><div class="l">Est. value / mo</div></div>'
                  f'<div class="card good"><div class="v">{_n(t.get("quick_wins"))}</div><div class="l">Quick-win keywords</div><div class="s">page-2 terms one push from page 1</div></div>'
                  f'</div>')
        mk = _kw_rows(seo["money_keywords"][:10], [("Money keyword","keyword","kw"),("Pos","position","pos"),
             ("Volume","search_volume","int"),("CPC","cpc","cpc"),("$ value/mo","money_value","money"),("Upside $","upside_value","money")]) if seo.get("money_keywords") else ""
        proof = ""
        if seo.get("proof_winning"):
            proof = ('<div class="lead" style="margin-top:16px">Proof the site <b>can</b> rank — already top-3 for these $ terms:</div>'
                     + _kw_rows(seo["proof_winning"][:5], [("Keyword","keyword","kw"),("Pos","position","pos"),("Volume","search_volume","int"),("$ value/mo","money_value","money")]))
        sec("seo", "Organic SEO health",
            "Where they rank today, and the high-value ‘money’ keywords sitting on page 2 (fastest ROI).", cards2 + mk + proof)

    # Keyword universe → funnel + product clusters
    uni = m.get("universe") or {}
    if uni.get("clusters"):
        fn = uni["funnel"]; tot = uni["totals"]
        def _tier(stage, cls, lab, w):
            f = fn.get(stage) or {}
            return (f'<div class="kf-tier {cls}" style="width:{w}%">'
                    f'<div class="kf-lab">{lab}</div>'
                    f'<div class="kf-m">{_n(f.get("keywords"))} keywords · {_n(f.get("volume"))} searches/mo · {_money(f.get("value"))}<br>'
                    f'<span style="opacity:.85">ranked {_n(f.get("ranked"))} · missing {_n(f.get("gap"))}</span></div></div>')
        funnel_html = ('<div class="kfunnel">'
                       + _tier("TOFU", "tofu", "TOFU · Awareness / research", 100)
                       + _tier("MOFU", "mofu", "MOFU · Consideration", 78)
                       + _tier("BOFU", "bofu", "BOFU · Ready to buy", 56) + '</div>')
        rows = []
        for cl in uni["clusters"][:10]:
            cov = round(100 * cl["ranked"] / max(1, cl["keywords"]))
            fm = cl["funnel"]; tt = max(1, fm["TOFU"] + fm["MOFU"] + fm["BOFU"])
            mix = (f'<span class="fmix" title="TOFU {fm["TOFU"]} · MOFU {fm["MOFU"]} · BOFU {fm["BOFU"]}">'
                   f'<i class="t" style="height:{max(3,round(16*fm["TOFU"]/tt))}px"></i>'
                   f'<i class="m" style="height:{max(3,round(16*fm["MOFU"]/tt))}px"></i>'
                   f'<i class="b" style="height:{max(3,round(16*fm["BOFU"]/tt))}px"></i></span>')
            tops = _esc(", ".join(x["keyword"] for x in cl["top"][:3]))
            rows.append(f'<tr><td class="l">{_esc(cl["label"])}<div style="font-weight:400;color:var(--faint);font-size:11px">{tops}</div></td>'
                        f'<td class="num">{_n(cl["keywords"])}</td><td class="num">{_n(cl["volume"])}</td><td class="num">{_money(cl["value"])}</td>'
                        f'<td class="num"><span class="covbar"><span style="width:{cov}%"></span></span>{cov}%</td>'
                        f'<td class="num">{mix}</td></tr>')
        cl_html = ('<div class="lead" style="margin-top:20px">Demand grouped by product / service — its size, value, how much they already rank for, and the funnel mix:</div>'
                   '<table class="cltable"><thead><tr><th class="l">Product / service cluster</th><th>Keywords</th>'
                   '<th>Searches/mo</th><th>$ value/mo</th><th>Coverage</th><th>Funnel mix</th></tr></thead>'
                   f'<tbody>{"".join(rows)}</tbody></table>'
                   '<div class="assump" style="margin-top:8px">Funnel mix bars: '
                   '<span class="fmix" style="vertical-align:middle"><i class="t" style="height:12px"></i></span> awareness · '
                   '<span class="fmix" style="vertical-align:middle"><i class="m" style="height:12px"></i></span> consideration · '
                   '<span class="fmix" style="vertical-align:middle"><i class="b" style="height:12px"></i></span> ready-to-buy.</div>')
        sec("universe", "Keyword universe & marketing funnel",
            f'The {_n(tot.get("keywords"))} money keywords across their category, mapped to the buyer funnel and clustered by product/service. '
            f'They rank for {_n(tot.get("ranked"))} and are missing {_n(tot.get("gap"))} — the gap is the growth map.',
            funnel_html + cl_html)

    # 4 Competitor scoreboard + gap
    if m.get("competitors"):
        nm = m["name"]
        # sort rivals biggest-first so the bar, "biggest rival" line and scoreboard all agree
        comps = sorted(m["competitors"][:5], key=lambda c: c.get("est_traffic") or 0, reverse=True)
        our_t = m.get("our_traffic") or 0
        total_t = our_t + sum((c.get("est_traffic") or 0) for c in comps)
        # share-of-voice bar
        sov_html = ""
        if total_t:
            palette = ["#1e3a8a", "#2563eb", "#3b82f6", "#60a5fa", "#93c5fd"]
            segs = [(nm, our_t, "#16a34a")] + [(c.get("domain"), c.get("est_traffic") or 0, palette[i % len(palette)]) for i, c in enumerate(comps)]
            other = max(0, total_t - sum(s[1] for s in segs))
            if other > 0:
                segs.append(("Others", other, "#cbd5e1"))
            bar = "".join(f'<div class="seg" style="flex:{max(1,s[1])};background:{s[2]}" title="{_esc(s[0])}: {_n(s[1])}/mo">'
                          f'{(str(round(100*s[1]/total_t))+"%") if s[1]/total_t>0.06 else ""}</div>' for s in segs)
            top_rival = comps[0]
            top_pct = round(100 * (top_rival.get("est_traffic") or 0) / total_t)
            sov_html = (f'<div class="lead" style="margin-top:0"><b>Share of organic search voice</b> — {_esc(nm)} holds ~'
                        f'{_n(m.get("sov"),1)}% of the estimated organic traffic in this competitive set, while their '
                        f'biggest rival <b>{_esc(top_rival.get("domain"))}</b> holds {top_pct}%. '
                        f'That imbalance is exactly the ground a focused organic program is built to win back.</div>'
                        f'<div class="sovbar">{bar}</div>')
        # head-to-head scoreboard
        def _scr(dom, traf, kws, avgpos, us=False):
            pct = round(100 * (traf or 0) / total_t) if total_t else 0
            return (f'<tr class="{"us" if us else ""}"><td class="l">{_esc(dom)}{" — you" if us else ""}</td>'
                    f'<td>{_n(kws)}</td><td>{_n(traf)}</td><td>{_pos(avgpos) if avgpos else "—"}</td><td>{pct}%</td></tr>')
        board = ('<table class="scoreboard"><thead><tr><th class="l">Domain</th><th>Organic keywords</th>'
                 '<th>Est. traffic/mo</th><th>Avg. position</th><th>Share</th></tr></thead><tbody>'
                 + _scr(nm, our_t, m["opportunity"].get("org_keywords"), None, True)
                 + "".join(_scr(c.get("domain"), c.get("est_traffic"), c.get("organic_keywords"), c.get("avg_position")) for c in comps)
                 + '</tbody></table>')
        # $-bleeding battle cards
        def _bleed(dom):
            o = [k for k in (m.get("outranked") or []) if k.get("competitor") == dom]
            g = [k for k in (m.get("keyword_gap") or []) if dom in (k.get("competitor_domains") or [])]
            return sum((k.get("cap_value") or 0) for k in o + g), len(o), len(g)
        cards_b = []
        for c in comps[:4]:
            val, no, ng = _bleed(c.get("domain"))
            if val <= 0 and no + ng == 0:
                continue
            cards_b.append(f'<div class="bcard"><div class="h">{_esc(c.get("domain"))}</div>'
                           f'<div class="m">Capturing an est. <b>{_money(val)}/mo</b> of search value across '
                           f'{no + ng} money keywords you lose to them — {ng} they rank for that you don’t, and {no} where they outrank you.</div></div>')
        battle_html = (f'<div class="lead" style="margin-top:18px">Where the revenue is bleeding — per rival:</div>'
                       f'<div class="battle">{"".join(cards_b)}</div>') if cards_b else ""
        # gap + outranked tables (fixed: volume + realistically CAPTURABLE value)
        gap = _kw_rows(m["keyword_gap"][:10], [("Keyword rivals win (you don’t rank)","keyword","kw"),
              ("Volume","volume","int"),("CPC","cpc","cpc"),("Capturable $/mo","cap_value","money")]) if m.get("keyword_gap") else ""
        outr = ('<div class="lead" style="margin-top:18px">Money keywords where a rival <b>outranks</b> you:</div>'
                + _kw_rows(m["outranked"][:8], [("Keyword","keyword","kw"),("Your pos","our_position","pos"),
                  ("Best rival pos","competitor_position","pos"),("Volume","volume","int"),("Capturable $/mo","cap_value","money")])) if m.get("outranked") else ""
        sec("competitor", "Competitor scoreboard & gap",
            "How they stack up against their real organic rivals — share of search, head-to-head, and exactly where money is leaking.",
            sov_html + board + battle_html + gap + outr)

    # 5 Content gap — demand-ranked topic cards the prospect has little/no content for (premium)
    if m.get("content_gap"):
        all_cg = [x for x in m["content_gap"] if isinstance(x, dict) and (x.get("topic") or x.get("keyword"))]
        cg_rows = sorted(all_cg, key=lambda x: x.get("total_volume") or 0, reverse=True)[:8]
        if cg_rows:
            maxv = max((x.get("total_volume") or 0) for x in cg_rows) or 1
            total_demand = sum((x.get("total_volume") or 0) for x in all_cg)
            rivals = sorted({x.get("domain") for x in all_cg if x.get("domain")})
            strip = ('<div class="cgstat">'
                     f'<div><div class="v">{_n(len(all_cg))}</div><div class="l">content themes missing or thin on their site</div></div>'
                     f'<div><div class="v">{_n(total_demand)}</div><div class="l">monthly searches they publish nothing for</div></div>'
                     f'<div><div class="v">{_n(len(rivals))}</div><div class="l">rival site{"s" if len(rivals) != 1 else ""} already ranking on this content</div></div></div>')

            def _clean_topic(t):
                # collapse adjacent repeated stems ("cable cover covers" → "cable cover") from DFS topic extraction
                out = []
                for w in (t or "").split():
                    if out and _stem(w.lower()) == _stem(out[-1].lower()):
                        continue
                    out.append(w)
                return " ".join(out)
            cards = []
            for x in cg_rows:
                topic = _clean_topic((x.get("topic") or x.get("keyword") or "").strip())
                vol = x.get("total_volume") or 0
                kc = x.get("keyword_count") or len(x.get("example_keywords") or [])
                dom = x.get("domain") or ""
                w = max(4, round(100 * vol / maxv))
                chips = "".join(f'<span class="cgchip">{_esc(k)}</span>' for k in (x.get("example_keywords") or [])[:4])
                cards.append(
                    f'<div class="cgcard"><div class="cgtop"><div class="cgt">{_esc(topic.title())}</div>'
                    f'<div class="cgv">{_n(vol)}<span>/mo</span></div></div>'
                    f'<div class="cgbar"><span style="width:{w}%"></span></div>'
                    f'<div class="cgmeta">{_n(kc)} related search{"es" if kc != 1 else ""}'
                    + (f' · owned by <b>{_esc(dom)}</b>' if dom else "") + "</div>"
                    + (f'<div class="cgchips">{chips}</div>' if chips else "") + "</div>")
            sec("content", "Content gap",
                "High-demand topics competitors publish for and rank on — and this business doesn’t. Each one is "
                "organic traffic, authority and trust being handed to a rival, month after month, for want of a page.",
                strip + '<div class="cggrid">' + "".join(cards) + "</div>")

    # 6 Backlinks / authority — render real metrics; if off-page analysis is off, show a clean note (not "available: False")
    bg = m.get("backlink_gap") or {}
    if bg:
        note = bg.get("note")
        metrics = {k: v for k, v in bg.items()
                   if k not in ("note", "available") and not isinstance(v, (list, dict, bool))}
        rows = "".join(f'<div class="comp"><div class="d">{_esc(k.replace("_", " ").title())}</div>'
                       f'<div class="m"><b>{_esc(v)}</b></div></div>' for k, v in metrics.items())
        body = rows
        if not rows and note:
            body = f'<div class="diag" style="border-left-color:var(--faint);background:var(--panel)">{_esc(note)}</div>'
        if body:
            sec("backlinks", "Backlinks & authority",
                "Off-page authority — the trusted external links that decide who ranks for the hardest, highest-value terms.",
                body)

    # 7 Technical / AI-search readiness
    geo = m.get("geo_aeo") or {}
    if geo and geo.get("checks"):
        sc = geo.get("score")
        cls = "" if sc is None else ("low" if sc < 45 else "mid" if sc < 70 else "")
        checks = "".join(
            f'<li class="{ "ok" if c.get("pass") else "no"}">{ "✓" if c.get("pass") else "✕"} {_esc(c.get("name"))}</li>'
            for c in (geo.get("checks") or [])[:12])
        body = (f'<div class="geo"><div class="score {cls}">{sc if sc is not None else "—"}</div>'
                f'<div><b>AI-search / answer-engine readiness</b><div class="lead" style="margin:2px 0 0">'
                f'How ready the site is to be surfaced by Google AI Overviews and AI assistants.</div></div></div>'
                f'<ul class="checks">{checks}</ul>')
        sec("technical", "Technical & AI-search readiness", "", body)

    # 7b The 90-day plan (quick wins → growth), from the already-built (previously unrendered) data
    qw = m.get("quick_wins") or []
    gp = m.get("growth_plan") or []
    def _pitem(x):
        eff = (x.get("effort") or "").lower(); imp = (x.get("impact") or "").lower()
        chips = ((f'<span class="chip hi">high impact</span>' if imp == "high" else f'<span class="chip md">{_esc(imp)} impact</span>' if imp else "")
                 + (f'<span class="chip lo">{_esc(eff)} effort</span>' if eff else ""))
        return f'<div class="pitem"><div class="t">{_esc(x.get("title"))}{chips}</div><div class="d">{_esc(x.get("detail"))}</div></div>'
    if qw or gp:
        p1 = "".join(_pitem(x) for x in qw[:5])
        p2 = "".join(_pitem(x) for x in gp[:5])
        plan = ''
        if p1:
            plan += f'<div class="phase"><div class="ph">Days 0–30 · Quick wins (fastest ROI)</div>{p1}</div>'
        if p2:
            plan += f'<div class="phase"><div class="ph">Days 30–90 · Growth plays (compounding)</div>{p2}</div>'
        sec("plan", "The 90-day plan",
            "Exactly what we would do first — sequenced quick-wins then the bigger growth plays, each ranked by impact and effort.", plan)

    # 8 Recommendation
    rec = m["recommendation"]
    steps = "".join(f'<div class="step"><div class="n">{i+1}</div><div><div class="h">{_esc(h)}</div><div class="d">{_esc(d)}</div></div></div>'
                    for i, (h, d) in enumerate(rec.get("steps") or []))
    up = f'<div class="headline" style="background:#0b1e46;margin-top:16px"><div class="big">{_money(upside)}/mo</div><div class="cap">estimated addressable search value we would work to capture</div></div>' if upside else ""
    sec("recommendation", "Our recommendation", "",
        f'<div class="rec"><div class="sum">{_esc(rec.get("summary"))}</div>{steps}</div>{up}')

    # Call to action + risk reversal (framed around what they're MISSING → why they need a team)
    sec("cta", "", "", (
        '<div class="cta"><h3>You’re leaving growth on the table</h3>'
        f'<p>This audit maps roughly <b>{rev_num("mo", r_mo)}/month</b> in new revenue {_esc(m["name"])} is currently '
        'handing to competitors and to gaps in its own site. Closing it means coordinated SEO, content, technical fixes and '
        'smarter paid — the kind of sustained, specialist work that’s hard to do in-house while running a business.</p>'
        '<a class="ctabtn" href="https://trafficradius.com.au/contact-us/" target="_blank" rel="noopener">Book a 30-minute strategy call →</a>'
        '<div class="rr">Australian team · transparent monthly reporting · we start with the fast wins so you see movement early.</div></div>'), num=False)

    # Methodology
    asmp = "".join(f'<li>{_esc(a)}</li>' for a in (m.get("assumptions") or []) if a)
    if asmp:
        sec("methodology", "Methodology & assumptions", "",
            f'<ul class="assump" style="margin:8px 0 0;padding-left:16px">{asmp}</ul>', num=False)

    # cover — the hero is the ticket-driven revenue opportunity (updates live everywhere)
    logo = m.get("_logo") or _logo_datauri()
    logo_chip = (f'<span class="logo"><img src="{logo}" alt="Traffic Radius"/></span>' if logo
                 else '<span class="logo">TR</span>')
    cover = (f'<div class="cover"><div class="brand">{logo_chip}<span>Traffic Radius</span></div>'
             f'<div class="kick">Digital Marketing Audit</div><h1>{_esc(m["name"])}</h1>'
             f'<div class="dom">{_esc(m["domain"])}</div>'
             f'<div class="meta">{meta_html}</div>'
             f'<div class="headline"><div class="big">{rev_num("mo", r_mo)}<span class="per">/ month</span></div>'
             f'<div class="cap">in new revenue they’re currently missing — quantified inside, at their own average sale value</div></div>'
             f'<div class="kick" style="margin-top:22px;color:#7ea6ff">Prepared {gen} · confidential</div></div>')

    footer = (f'<div class="footer"><b>Traffic Radius</b> — Digital Marketing Audit for {_esc(m["name"])} ({_esc(m["domain"])}). '
              f'Generated {gen}. SEO &amp; keyword analysis by Traffic Radius; ad activity independently verified via '
              f'Google’s public Ads Transparency Center; traffic &amp; value figures '
              f'are estimates (see methodology). Prepared for a discovery conversation, not a guarantee of results.</div>')

    # LOGICAL ORDER: diagnosis → the money → evidence (ads/seo/keywords/competitors/content/tech) →
    #                the plan → recommendation → act. Numbered at emit time.
    ORDER = ["health", "exec", "revenue", "gads", "seo", "universe", "competitor",
             "content", "backlinks", "technical", "plan", "recommendation", "cta", "methodology"]
    n = 0
    secs = []
    for k in ORDER:
        if k not in S:
            continue
        title, lead, bd, num = S[k]
        head = ""
        if title:
            if num:
                n += 1
            tag = f'<span class="secnum">{n}</span>' if num else ""
            lead_h = f'<div class="lead">{_esc(lead)}</div>' if lead else ""
            head = f'<h2>{tag}{_esc(title)}</h2>{lead_h}'
        secs.append(f'<div class="sec">{head}{bd}</div>')

    js = ('<script>function recalcRev(el){var r=el.closest(".page")||document;'
          'var t=+r.dataset.traf||0,v=+r.dataset.v2l||0.03,l=+r.dataset.l2s||0.25,k=+el.value||0;'
          'var leads=t*v,sales=leads*l,mo=Math.round(sales*k),yr=mo*12,q=mo*3;'
          'var f=function(x,d){return x.toLocaleString(undefined,{maximumFractionDigits:d||0,minimumFractionDigits:d||0});};'
          'r.querySelectorAll("[data-rev]").forEach(function(n){var kd=n.dataset.rev,'
          'val=kd=="mo"?mo:kd=="yr"?yr:kd=="q"?q:kd=="sales"?sales:kd=="leads"?leads:kd=="ticket"?k:0,'
          'money=(kd=="mo"||kd=="yr"||kd=="q"||kd=="ticket"),dec=(kd=="sales"||kd=="leads")?1:0;'
          'n.textContent=(money?"$":"")+f(val,dec);});}'
          'document.addEventListener("DOMContentLoaded",function(){document.querySelectorAll(".tkt-input").forEach(function(i){recalcRev(i);});});</script>')
    pd = f'data-traf="{r_traf}" data-v2l="{r_v2l}" data-l2s="{r_l2s}"'
    body_html = cover + "".join(secs) + footer + js
    if not standalone:
        return f'<style>{_AUDIT_CSS}</style><div class="page" {pd}>{body_html}</div>'
    return (f'<!doctype html><html><head><meta charset="utf-8"><title>Audit — {_esc(m["name"])}</title>'
            f'<style>{_AUDIT_CSS}</style></head><body><div class="page" {pd}>{body_html}</div></body></html>')


# ============================================================================
# KEYWORD UNIVERSE → product clusters → marketing-funnel mapping
# ============================================================================
import re as _re
from collections import Counter, defaultdict

# funnel-stage intent lexicon (research-based mapping of search intent → funnel stage)
_BOFU = (" buy "," price "," prices "," pricing "," cost "," costs "," quote "," quotes "," cheap ",
         " sale "," for sale "," supplier "," suppliers "," supplies "," near me "," installation ",
         " install "," installer "," hire "," rental "," repair "," service "," services "," company ",
         " companies "," contractor "," order "," shop "," store "," wholesale "," custom "," bespoke ",
         " manufacturer "," specialist "," fitting "," fitted "," replacement "," quote ")
_MOFU = (" best "," top "," review "," reviews "," vs "," versus "," compare "," comparison "," types ",
         " type of "," brands "," brand "," alternative "," alternatives "," which "," rated "," recommended ")
_TOFU = (" how "," how to "," what "," what is "," why "," guide "," ideas "," idea "," tips "," diy ",
         " meaning "," benefits "," examples "," explained "," vs "," difference ")
_STOP = {"the","and","for","with","your","you","our","are","can","get","from","that","this","near","me",
         "in","on","of","to","a","an","at","by","or","is","best","buy","how","what","why","top","cheap",
         "price","cost","sale","new","used","free","australia","australian","au","com","www","service","services"}
_AU_LOC = {"sydney","melbourne","brisbane","perth","adelaide","canberra","hobart","darwin","gold","coast",
           "newcastle","wollongong","geelong","nsw","vic","qld","wa","sa","tas","act","nt","australia"}


# a product term carrying a size / spec / model number ("100mm pvc pipe", "dn300", "pn16", "class 12")
# is a ready-to-buy query — the searcher already knows exactly what they want.
_SPEC_RE = _re.compile(r"\b\d{2,}\s?(?:mm|cm|nb|dn|pn|inch|in|kg|mpa|bar|class|kva|kw|ltr|litre|amp)\b"
                       r"|\bdn\s?\d+|\bpn\s?\d+|\bclass\s?\d+", _re.I)


def _assign_stages(pool: list) -> None:
    """Map every keyword in the pool to a marketing-funnel stage (TOFU/MOFU/BOFU) in place.

    Intent is genuinely fuzzy for bare product terms, so we combine an explicit-intent lexicon with
    CPC as a commercial-intent proxy (advertisers pay more per click the closer a term is to purchase).
    CPC bands are computed as TERCILES *within this prospect's own keyword set*, so the split adapts to
    the niche and always yields a realistic spread instead of dumping everything into one stage."""
    cpcs = sorted(r["cpc"] for r in pool if (r.get("cpc") or 0) > 0)

    def _pct(p):
        if not cpcs:
            return 0.0
        return cpcs[min(len(cpcs) - 1, int(p * len(cpcs)))]

    lo, hi = _pct(0.40), _pct(0.75)   # bottom 40% CPC = TOFU band; top 25% = BOFU band
    for r in pool:
        kw = r.get("keyword") or ""
        k = " " + kw.lower() + " "
        cpc = r.get("cpc") or 0
        nwords = len(_re.findall(r"[a-z0-9]+", kw.lower()))
        # 1. explicit intent words always win
        if any(t in k for t in _TOFU):
            r["stage"] = "TOFU"; continue
        if any(t in k for t in _BOFU) or _SPEC_RE.search(kw):
            r["stage"] = "BOFU"; continue
        if any(t in k for t in _MOFU):
            r["stage"] = "MOFU"; continue
        # 2. otherwise infer from commercial intent (CPC) refined by specificity (longer = more precise)
        if cpc and hi > 0 and cpc >= hi:
            r["stage"] = "BOFU"
        elif nwords >= 3 and cpc and cpc >= lo:
            r["stage"] = "BOFU"                    # specific multi-word product term = ready to buy
        elif nwords <= 2 and cpc <= lo:
            r["stage"] = "TOFU"                    # broad, low-CPC head term = awareness / research
        else:
            r["stage"] = "MOFU"                    # mid-CPC category term = consideration


def _stem(t: str) -> str:
    """Cheap singularisation so 'pipe'/'pipes' and 'valve'/'valves' cluster together."""
    if len(t) > 4 and t.endswith("s") and not t.endswith("ss"):
        return t[:-3] if t.endswith("ies") else t[:-1]
    return t


def _sig_tokens(kw: str) -> list:
    toks = _re.findall(r"[a-z][a-z]+", (kw or "").lower())
    return [_stem(t) for t in toks if len(t) > 2 and t not in _STOP and t not in _AU_LOC]


def build_keyword_universe(ranked: list, gap: list, *, brands=None) -> dict:
    """Assemble a keyword UNIVERSE from the terms the prospect ranks for + the competitor gap, then
    (1) map every keyword to a marketing-funnel stage by search intent, and (2) cluster keywords into
    product/service themes with coverage (ranked vs missing) — the opportunity map."""
    brands = set(brands or set())
    try:
        from .enrichment.dataforseo import is_branded as _is_branded
    except Exception:
        def _is_branded(kw, b):
            return False
    pool = []
    for k in (ranked or []):
        vol = k.get("search_volume") or k.get("volume") or 0
        cpc = k.get("cpc") or 0
        pos = k.get("position")
        if vol < 20 or _is_branded(k.get("keyword"), brands):   # drop the prospect's own brand terms
            continue
        pool.append({"keyword": k.get("keyword"), "vol": vol, "cpc": cpc, "pos": pos,
                     "value": round(vol * cpc), "ranked": True,
                     "traffic": round(vol * _ctr(pos or 999)), "stage": None})
    seen = {r["keyword"] for r in pool}
    for k in (gap or []):
        kw = k.get("keyword")
        if not kw or kw in seen or _is_branded(kw, brands):
            continue
        vol = k.get("volume") or k.get("search_volume") or 0
        cpc = k.get("cpc") or 0
        if vol < 20:
            continue
        pool.append({"keyword": kw, "vol": vol, "cpc": cpc, "pos": None,
                     "value": k.get("money_value") or round(vol * cpc), "ranked": False,
                     "traffic": k.get("est_capture_traffic") or round(vol * _ctr(5)),
                     "stage": None})
    # ---- map every keyword to a funnel stage (adaptive: intent lexicon + CPC terciles) ----
    _assign_stages(pool)
    # ---- cluster by dominant product token ----
    freq = Counter()
    for r in pool:
        r["_st"] = _sig_tokens(r["keyword"])
        for t in set(r["_st"]):
            freq[t] += 1
    # tokens that appear in >33% of the universe are too generic to be a cluster theme (e.g. "pipe" for a
    # pipe maker) — cluster on the most DISTINCTIVE shared token instead, so one head term doesn't swallow
    # everything. Fall back to the generic head only when a keyword has no more specific token.
    n_pool = len(pool)
    generic = {t for t, c in freq.items() if n_pool and c > 0.33 * n_pool}
    clusters = defaultdict(list)
    for r in pool:
        st = r["_st"]
        specific = [t for t in st if t not in generic]
        cand = specific or st
        key = max(cand, key=lambda t: (freq[t], len(t))) if cand else "other"
        clusters[key].append(r)
    # merge singletons / tiny into a long-tail bucket
    cl_out = []
    longtail = []
    for key, rows in clusters.items():
        if len(rows) < 2 or key == "other":
            longtail += rows
            continue
        cl_out.append((key, rows))

    def _agg(rows):
        stages = Counter(r["stage"] for r in rows)
        return {
            "keywords": len(rows),
            "volume": sum(r["vol"] for r in rows),
            "value": sum(r["value"] for r in rows),
            "ranked": sum(1 for r in rows if r["ranked"]),
            "gap": sum(1 for r in rows if not r["ranked"]),
            "traffic": sum(r["traffic"] for r in rows if not r["ranked"]),   # addressable (not-yet-ranked)
            "funnel": {"TOFU": stages["TOFU"], "MOFU": stages["MOFU"], "BOFU": stages["BOFU"]},
            "top": sorted(rows, key=lambda r: r["value"], reverse=True)[:4],
        }
    clusters_final = []
    for key, rows in cl_out:
        a = _agg(rows)
        a["label"] = " / ".join(k.title() for k in _cluster_label(rows, key))
        clusters_final.append(a)
    clusters_final.sort(key=lambda a: a["value"], reverse=True)
    if longtail:
        a = _agg(longtail); a["label"] = "Long-tail / other"; clusters_final.append(a)

    # ---- funnel roll-up across the whole universe ----
    def _stage_agg(stage):
        rows = [r for r in pool if r["stage"] == stage]
        return {"keywords": len(rows), "volume": sum(r["vol"] for r in rows),
                "value": sum(r["value"] for r in rows),
                "ranked": sum(1 for r in rows if r["ranked"]),
                "gap": sum(1 for r in rows if not r["ranked"]),
                "traffic": sum(r["traffic"] for r in rows if not r["ranked"])}
    funnel = {s: _stage_agg(s) for s in ("TOFU", "MOFU", "BOFU")}
    return {
        "totals": {"keywords": len(pool), "volume": sum(r["vol"] for r in pool),
                   "value": sum(r["value"] for r in pool),
                   "ranked": sum(1 for r in pool if r["ranked"]),
                   "gap": sum(1 for r in pool if not r["ranked"]),
                   "addressable_traffic": sum(r["traffic"] for r in pool if not r["ranked"])},
        "funnel": funnel,
        "clusters": clusters_final[:10],
    }


def _cluster_label(rows, key):
    """A readable 1–2 word label for a cluster: its most common significant word(s), rendered in their
    original surface form (stem for counting so pipe/pipes merge, but display the real word so proper
    nouns like 'Paris' aren't mangled to 'Pari')."""
    words = Counter()
    surface = {}
    for r in rows:
        for w in _re.findall(r"[a-z][a-z]+", (r["keyword"] or "").lower()):
            if len(w) > 2 and w not in _STOP and w not in _AU_LOC:
                s = _stem(w)
                words[s] += 1
                surface.setdefault(s, w)
    common = [w for w, _ in words.most_common(2)]
    return [surface.get(w, w) for w in common] or [surface.get(_stem(key), key)]


# ============================================================================
# health scoring, diagnosis, share-of-voice, SVG gauges (audit-improve panel)
# ============================================================================
def _health_scores(est_org_value, seo_tot, ads_m, sov, geo, n_comp) -> dict:
    """A 0–100 digital-health scorecard (higher = healthier). Illustrative composite from the signals
    we hold — its job is to show at a glance where the prospect is weak (i.e. where the opportunity is)."""
    t = seo_tot or {}
    kw = t.get("keywords") or 0
    win = t.get("winning") or 0
    quick = t.get("quick_wins") or 0
    seo = max(8, min(92, round(win * 2.4 + min(45, kw / 6) + min(14, quick))))
    ads = 86 if (ads_m or {}).get("running") else 18
    comp = round(min(95, max(6, (sov or 0) * 2.2))) if sov is not None else 45
    tech = geo.get("score") if geo and geo.get("score") is not None else 45
    overall = round(0.40 * seo + 0.15 * ads + 0.25 * comp + 0.20 * tech)
    return {"overall": overall, "seo": seo, "ads": ads, "competitive": comp, "technical": int(tech)}


def _diagnosis(name, running_ads, ads_m, quickwin_value, gap_capturable, sov, geo, comps):
    """One tailored paragraph that names the prospect's specific situation — the emotional hook."""
    a = []
    if running_ads and (ads_m or {}).get("years_active"):
        a.append(f"{name} has invested in Google Ads for {_n(ads_m['years_active'])} years — the intent to grow is clearly there")
    elif running_ads:
        a.append(f"{name} is actively running Google Ads — the intent to grow is there")
    else:
        a.append(f"{name} has real organic foundations to build on")
    if comps:
        a.append(f"yet in the organic results — where the free, compounding traffic lives — rivals like {comps[0].get('domain')} are pulling ahead")
    if sov is not None and sov < 25:
        a.append(f"they currently hold only ~{_n(sov,1)}% of their category's organic search visibility")
    total = (quickwin_value or 0) + (gap_capturable or 0)
    tail = (f" The result is roughly {_money(total)}/mo of realistically capturable search value going uncaptured — "
            "money their ad budget is effectively subsidising." if total else
            " That leaves a clear, capturable gap between their ad spend and their organic return.")
    return ", ".join(a) + "." + tail


def _gauge(score, label, size=92):
    """A CSS conic-gradient gauge (PDF-safe): coloured arc + centre number."""
    s = max(0, min(100, int(score or 0)))
    col = "#dc2626" if s < 40 else ("#d97706" if s < 65 else "#16a34a")
    return (f'<div class="gauge"><div class="ring" style="width:{size}px;height:{size}px;'
            f'background:conic-gradient({col} {s * 3.6:.0f}deg,#e2e8f0 0)">'
            f'<div class="hole"><span style="color:{col}">{s}</span></div></div>'
            f'<div class="glab">{_esc(label)}</div></div>')


# ============================================================================
# Export renderers — PDF (WeasyPrint, pixel-faithful to the on-screen report)
# and DOCX (python-docx, a clean editable version built from the model).
# ============================================================================
def render_audit_pdf(model: dict) -> bytes:
    """Render the audit to PDF bytes via WeasyPrint from the same self-contained HTML (embedded
    data-URI ad images + logo). Ships with the app; the Docker image carries Pango/Cairo."""
    from weasyprint import HTML
    html = render_audit_html(model, standalone=True)
    return HTML(string=html).write_pdf()


def _docx_shade(cell, hex_no_hash: str) -> None:
    """Set a table cell background (python-docx has no direct API for it)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_no_hash)
    tcPr.append(shd)


def render_audit_docx(model: dict) -> bytes:
    """Build a clean, editable DOCX of the audit from the model (not the HTML). Not pixel-identical
    to the PDF — it's the professional editable version: headings, KPI lines, and data tables."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    NAVY = RGBColor(0x0B, 0x1E, 0x46)
    BLUE = RGBColor(0x25, 0x63, 0xEB)
    GREEN = RGBColor(0x16, 0xA3, 0x4A)
    AMBER = RGBColor(0xD9, 0x77, 0x06)
    RED = RGBColor(0xDC, 0x26, 0x26)
    MUTED = RGBColor(0x64, 0x74, 0x8B)
    INK = RGBColor(0x0F, 0x17, 0x2A)

    m = model
    b = m.get("business") or {}
    op = m.get("opportunity") or {}
    ads = m.get("ads") or {}
    seo = m.get("seo") or {}
    rev = m.get("revenue") or {}
    uni = m.get("universe") or {}
    gen = (m.get("generated") or "")[:10] or date.today().isoformat()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    def para(text="", *, bold=False, italic=False, size=10.5, color=None, align=None, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if align is not None:
            p.alignment = align
        if text:
            r = p.add_run(text)
            r.bold = bold; r.italic = italic; r.font.size = Pt(size)
            r.font.color.rgb = color or INK
        return p

    _secno = [0]
    def section(title, lead=None, *, number=True):
        _secno[0] += 1 if number else 0
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(2)
        if number:
            n = p.add_run(f"{_secno[0]}  "); n.bold = True; n.font.size = Pt(15); n.font.color.rgb = BLUE
        t = p.add_run(title); t.bold = True; t.font.size = Pt(15); t.font.color.rgb = NAVY
        if lead:
            para(lead, italic=True, size=10, color=MUTED, space_after=8)

    def kpi_row(items):
        """items = [(value, label), ...] as a light table row of KPI cards."""
        if not items:
            return
        tbl = doc.add_table(rows=2, cols=len(items))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (val, lab) in enumerate(items):
            c0 = tbl.cell(0, i); c1 = tbl.cell(1, i)
            _docx_shade(c0, "F1F5FB"); _docx_shade(c1, "F1F5FB")
            r = c0.paragraphs[0].add_run(str(val)); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
            r2 = c1.paragraphs[0].add_run(str(lab)); r2.font.size = Pt(8.5); r2.font.color.rgb = MUTED
        para(space_after=4)

    def data_table(headers, rows, *, aligns=None):
        if not rows:
            return
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            run = cell.paragraphs[0].add_run(str(h)); run.bold = True; run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for row in rows:
            cells = tbl.add_row().cells
            for i, v in enumerate(row):
                r = cells[i].paragraphs[0].add_run("" if v is None else str(v))
                r.font.size = Pt(9)
                if aligns and i < len(aligns) and aligns[i] == "r":
                    cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para(space_after=6)

    # ---- COVER ----
    para("TRAFFIC RADIUS", bold=True, size=11, color=BLUE, space_after=0)
    para("Digital Marketing Audit", size=9, color=MUTED, space_after=10)
    h = doc.add_paragraph(); h.paragraph_format.space_after = Pt(2)
    rr = h.add_run(m.get("name") or m.get("domain") or ""); rr.bold = True; rr.font.size = Pt(24); rr.font.color.rgb = NAVY
    para(m.get("domain") or "", size=11, color=BLUE, space_after=8)
    meta = []
    if b.get("industry"): meta.append(f"Industry: {b['industry']}")
    if b.get("location"): meta.append(f"Location: {b['location']}")
    if b.get("revenue_musd"): meta.append(f"Revenue: ${_n(b['revenue_musd'])}M")
    if op.get("org_keywords"): meta.append(f"Organic keywords: {_n(op['org_keywords'])}")
    if meta:
        para(" · ".join(meta), size=9.5, color=MUTED, space_after=10)
    r_mo = rev.get("monthly")
    if r_mo:
        hp = doc.add_paragraph(); hp.paragraph_format.space_after = Pt(2)
        big = hp.add_run(f"{_money(r_mo)}/month"); big.bold = True; big.font.size = Pt(20); big.font.color.rgb = GREEN
        para("in new revenue they're currently missing — quantified inside, at their own average sale value.",
             italic=True, size=9.5, color=MUTED)
    para(f"Prepared {gen} · confidential", size=8.5, color=MUTED, space_after=2)

    # ---- HEALTH SCORECARD ----
    hlth = m.get("health") or {}
    if hlth:
        ov = hlth.get("overall") or 0
        grade = "A" if ov >= 80 else "B" if ov >= 65 else "C" if ov >= 50 else "D" if ov >= 35 else "F"
        section("Digital health scorecard",
                "An honest, at-a-glance grade of the business online today — the red gaps are the ground competitors are taking.")
        para(f"Overall grade: {grade}  ({ov}/100)", bold=True, size=13,
             color=GREEN if ov >= 65 else AMBER if ov >= 45 else RED)
        DIMS = [("seo", "Organic visibility"), ("ads", "Paid presence"),
                ("competitive", "Competitive position"), ("technical", "AI-search readiness")]
        data_table(["Dimension", "Score / 100", "Target"],
                   [[lab, int(hlth.get(k) or 0), 80] for k, lab in DIMS], aligns=["l", "r", "r"])
        if m.get("diagnosis"):
            para(m["diagnosis"], size=10)

    # ---- EXEC SUMMARY ----
    if m.get("findings"):
        section("Executive summary", "The headline findings and the size of the opportunity.")
        kpis = []
        if r_mo: kpis.append((f"{_money(r_mo)}", "New revenue / mo"))
        if op.get("est_org_value") is not None: kpis.append((_money(op["est_org_value"]), "Organic value / mo"))
        if ads.get("running"): kpis.append((str(ads.get("count") or "Yes"), "Live Google Ads"))
        if op.get("org_keywords"): kpis.append((_n(op["org_keywords"]), "Ranked keywords"))
        kpi_row(kpis[:4])
        for f in m["findings"]:
            para(f"•  {f.get('text','')}", size=10, space_after=4)

    # ---- REVENUE OPPORTUNITY ----
    if rev.get("addressable_traffic"):
        section("The revenue opportunity",
                "What the missed search traffic is worth in the prospect's own revenue terms.")
        para(f"Average sale / project value used: {_money(rev.get('avg_ticket'))}", bold=True, size=10)
        data_table(["Extra visits/mo", "Enquiries/mo", "New sales/mo", "New revenue/mo", "Per year"],
                   [[_n(rev.get("addressable_traffic")), _n(rev.get("leads_per_mo"), 1),
                     _n(rev.get("sales_per_mo"), 1), _money(r_mo), _money((r_mo or 0) * 12)]],
                   aligns=["r", "r", "r", "r", "r"])
        para(f"Assumes {int((rev.get('visitor_to_lead') or .03)*100)}% visitor→enquiry and "
             f"{int((rev.get('lead_to_sale') or .25)*100)}% enquiry→sale (conservative service-business rates).",
             italic=True, size=8.5, color=MUTED)

    # ---- GOOGLE ADS ----
    if ads.get("running") or ads.get("count"):
        section("Google Ads audit",
                "Independently verified from Google's public Ads Transparency Center.")
        kpi_row([("✓ Confirmed" if ads.get("running") else "—", "Running Google Ads"),
                 (str(ads.get("count") or "—"), "Live creatives"),
                 (str(ads.get("since") or "—")[:4] if ads.get("since") else "—", "Advertising since"),
                 ("Verified" if ads.get("verified") else "—", "Advertiser identity")])
        if ads.get("formats"):
            para("Formats: " + " · ".join(f"{k}: {v}" for k, v in ads["formats"].items()), size=9.5, color=MUTED)

    # ---- SEO HEALTH ----
    t = seo.get("totals") or {}
    if seo.get("money_keywords") or t:
        section("Organic SEO health",
                "Where they rank today, and the high-value 'money' keywords sitting on page 2 (fastest ROI).")
        kpi_row([(_n(t.get("keywords") or seo.get("keyword_count_total")), "Ranked keywords"),
                 (_n(t.get("est_organic_traffic")), "Est. visits / mo"),
                 (_money(t.get("est_traffic_value")), "Est. value / mo"),
                 (_n(t.get("quick_wins")), "Quick-win keywords")])
        mk = seo.get("money_keywords") or []
        if mk:
            data_table(["Money keyword", "Pos", "Volume", "CPC", "$ value/mo", "Upside $"],
                       [[k.get("keyword"), _pos(k.get("position")), _n(k.get("search_volume")),
                         _money(k.get("cpc"), 2), _money(k.get("money_value")), _money(k.get("upside_value"))]
                        for k in mk[:12]], aligns=["l", "r", "r", "r", "r", "r"])

    # ---- KEYWORD UNIVERSE ----
    if uni.get("clusters"):
        fn = uni.get("funnel") or {}
        section("Keyword universe & marketing funnel",
                "The money keywords mapped to the buyer funnel and clustered by product / service.")
        data_table(["Funnel stage", "Keywords", "Searches/mo", "$ value/mo", "Ranked", "Missing"],
                   [[lab, _n((fn.get(s) or {}).get("keywords")), _n((fn.get(s) or {}).get("volume")),
                     _money((fn.get(s) or {}).get("value")), _n((fn.get(s) or {}).get("ranked")),
                     _n((fn.get(s) or {}).get("gap"))]
                    for s, lab in [("TOFU", "TOFU · Awareness"), ("MOFU", "MOFU · Consideration"), ("BOFU", "BOFU · Ready to buy")]],
                   aligns=["l", "r", "r", "r", "r", "r"])
        data_table(["Product / service cluster", "Keywords", "Searches/mo", "$ value/mo", "Coverage"],
                   [[cl.get("label"), _n(cl.get("keywords")), _n(cl.get("volume")), _money(cl.get("value")),
                     f"{round(100*(cl.get('ranked') or 0)/max(1, cl.get('keywords') or 1))}%"]
                    for cl in uni["clusters"][:10]], aligns=["l", "r", "r", "r", "r"])

    # ---- COMPETITOR ----
    if m.get("competitors"):
        comps = sorted(m["competitors"][:5], key=lambda c: c.get("est_traffic") or 0, reverse=True)
        section("Competitor scoreboard & gap",
                "How they stack up against their real organic rivals — and exactly where money is leaking.")
        if m.get("sov") is not None:
            para(f"Share of organic search voice: ~{_n(m.get('sov'), 1)}% "
                 f"(biggest rival {comps[0].get('domain')} holds "
                 f"{round(100*(comps[0].get('est_traffic') or 0)/max(1,(m.get('our_traffic') or 0)+sum((c.get('est_traffic') or 0) for c in comps)))}%).",
                 size=10)
        data_table(["Domain", "Organic keywords", "Est. traffic/mo", "Avg. position"],
                   [[f"{m.get('name')} (you)", _n(op.get("org_keywords")), _n(m.get("our_traffic")), "—"]]
                   + [[c.get("domain"), _n(c.get("organic_keywords")), _n(c.get("est_traffic")),
                       _pos(c.get("avg_position")) if c.get("avg_position") else "—"] for c in comps],
                   aligns=["l", "r", "r", "r"])
        gap = m.get("keyword_gap") or []
        if gap:
            para("Keywords rivals win that you don't rank for:", bold=True, size=10, space_after=2)
            data_table(["Keyword", "Volume", "CPC", "Capturable $/mo"],
                       [[k.get("keyword"), _n(k.get("volume")), _money(k.get("cpc"), 2), _money(k.get("cap_value"))]
                        for k in gap[:10]], aligns=["l", "r", "r", "r"])

    # ---- CONTENT GAP ----
    cg = [x for x in (m.get("content_gap") or []) if isinstance(x, dict)]
    if cg:
        cg = sorted(cg, key=lambda x: x.get("total_volume") or 0, reverse=True)[:10]
        section("Content gap",
                "High-demand topics competitors publish for and rank on — and this business doesn't.")
        data_table(["Topic", "Searches/mo", "Related", "Covered by"],
                   [[(x.get("topic") or x.get("keyword") or "").title(), _n(x.get("total_volume")),
                     _n(x.get("keyword_count")), x.get("domain") or ""] for x in cg],
                   aligns=["l", "r", "r", "l"])

    # ---- TECHNICAL / AI READINESS ----
    geo = m.get("geo_aeo") or {}
    if geo and geo.get("checks"):
        section("Technical & AI-search readiness",
                "Whether AI & answer engines can find, trust and surface the site.")
        if geo.get("score") is not None:
            para(f"AI-search readiness score: {geo.get('score')}/100", bold=True, size=11,
                 color=GREEN if geo["score"] >= 70 else AMBER if geo["score"] >= 45 else RED)
        for c in (geo.get("checks") or [])[:12]:
            ok = c.get("pass")
            para(f"{'✓' if ok else '✗'}  {c.get('label') or c.get('name') or ''}", size=9.5,
                 color=GREEN if ok else RED, space_after=2)

    # ---- 90-DAY PLAN / RECOMMENDATION ----
    rec = m.get("recommendation") or {}
    if rec.get("steps"):
        section("The opportunity & our recommendation", rec.get("summary"))
        for i, step in enumerate(rec["steps"], 1):
            title, detail = (step if isinstance(step, (list, tuple)) else (step, ""))
            para(f"{i}. {title}", bold=True, size=10.5, color=NAVY, space_after=1)
            if detail:
                para(detail, size=9.5, color=INK, space_after=5)
        if rec.get("total_upside"):
            para(f"Estimated addressable search upside: {_money(rec['total_upside'])}/mo.",
                 bold=True, size=10.5, color=GREEN)

    # ---- METHODOLOGY ----
    asmp = [a for a in (m.get("assumptions") or []) if a]
    if asmp:
        section("Methodology & assumptions", number=False)
        for a in asmp:
            para(f"•  {a}", size=8.5, color=MUTED, space_after=3)

    para()
    para(f"Traffic Radius — Digital Marketing Audit for {m.get('name')} ({m.get('domain')}). "
         f"Generated {gen}. SEO & keyword analysis by Traffic Radius; ad activity independently verified via "
         f"Google's public Ads Transparency Center; "
         f"traffic & value figures are estimates. Prepared for a discovery conversation, not a guarantee of results.",
         italic=True, size=8, color=MUTED)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
